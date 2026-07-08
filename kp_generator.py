"""
Генератор технико-коммерческого предложения (КП) в формате docx.
Использует шаблон СибКомплект и подставляет данные проекта.
"""

import os
import copy
import datetime
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Путь к шаблону (образец КП)
TEMPLATE_PATH = Path(__file__).parent / "kp_template.docx"

# Месяцы для форматирования даты
MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
}


def _format_price(val: float) -> str:
    """Форматирует цену как '1 049 911,00'"""
    try:
        return f"{float(val):,.2f}".replace(",", " ").replace(".", ",")
    except Exception:
        return "0,00"


def _set_cell_text(cell, text, bold=False, size=9, align="left"):
    """Устанавливает текст ячейки с форматированием"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"


def _set_name_cell(cell, text, size=9):
    """
    Записывает наименование позиции со списком состава.
    Вводная часть остаётся абзацем, а перечисление через '; ' разбивается
    на строки-пункты с маркером.
    """
    cell.text = ""
    text = str(text).strip()

    # Разделяем на вводную часть и состав по маркеру "в составе:" / "в составе ячеек" и т.п.
    intro = text
    items_part = ""
    for sep in ["в составе:", "в составе ", ", в составе", "в состав входит:", ": "]:
        idx = text.find(sep)
        if idx != -1:
            intro = text[:idx + len(sep)].rstrip()
            items_part = text[idx + len(sep):].strip()
            break

    first_para = cell.paragraphs[0]

    if not items_part:
        # Нет списка — пишем как обычный текст
        run = first_para.add_run(intro)
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        return

    # Вводная часть
    run = first_para.add_run(intro)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"

    # Разбиваем состав по "; " на пункты
    items = [it.strip().rstrip(";.") for it in items_part.split(";") if it.strip()]
    for it in items:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Pt(8)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        r = para.add_run(f"– {it}")
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"


def _replace_in_paragraph(paragraph, replacements):
    """Заменяет текст в параграфе, сохраняя форматирование первого run"""
    full_text = paragraph.text
    new_text = full_text
    for key, val in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, str(val))

    if new_text != full_text:
        # Сохраняем форматирование первого run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            font_name = first_run.font.name
            # Очищаем все runs
            for run in paragraph.runs:
                run.text = ""
            # Пишем новый текст в первый run
            paragraph.runs[0].text = new_text
            paragraph.runs[0].font.size = font_size
            paragraph.runs[0].font.bold = font_bold
            paragraph.runs[0].font.name = font_name


def generate_kp(
    output_path: str,
    object_name: str,
    equipment_title: str,
    positions: list[dict],
    delivery_address: str = "",
    project_docs: str = "",
    ish_number: str = None,
    validity_days: int = 10,
) -> str:
    """
    Генерирует КП по шаблону.

    positions: [{"name": "...", "price": 1049911.00, "qty": 1}, ...]
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Шаблон КП не найден: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))

    now = datetime.datetime.now()
    date_str = f"«{now.day:02d}» {MONTHS_RU[now.month]} {now.year} г."
    validity_date = (now + datetime.timedelta(days=validity_days)).strftime("%d.%m.%Y")

    if not ish_number:
        ish_number = f"{now.strftime('%y%m')}-{now.strftime('%d')}"

    # Заменяемые переменные в тексте
    replacements = {
        "2506-15": ish_number,
        "«04» июня 2025 г.": date_str,
        "ПАО «Россети Сибирь» - «Алтайэнерго» Объект: Южный тракт 18. ООО «СЗ «Адалин»": object_name or "Объект",
        "Комплектная трансформаторная подстанция внутреннего исполнения 2КТПВ-СК-КК-1250кВА-6/0,4кВ УХЛ1": equipment_title,
        "г. Барнаул, Южный тракт, 18": delivery_address or "уточняется",
        "16.06.2025": validity_date,
    }

    # Заменяем docs в тексте если указаны
    if project_docs:
        replacements["(12-2022-ЭМО.ТП, 12-2022- НЭС.ТП)"] = f"({project_docs})"

    # Проходим по всем параграфам и заменяем
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Заполняем таблицу
    table = doc.tables[0]

    # Строка 0 — шапка, строка 1 — образец данных (5 ячеек), последняя — итого (2 ячейки)
    header_row = table.rows[0]
    sample_data_row = table.rows[1]  # берём как образец для копирования
    total_row = table.rows[-1]

    # Запоминаем XML образца строки данных ДО удаления
    sample_row_xml = copy.deepcopy(sample_data_row._element)

    # Удаляем все строки данных (между шапкой и итогом)
    rows_to_remove = list(table.rows)[1:-1]
    for row in rows_to_remove:
        row._element.getparent().remove(row._element)

    # Добавляем новые строки с позициями перед итоговой
    from docx.table import _Row
    grand_total = 0.0
    for idx, pos in enumerate(positions, 1):
        price = float(pos.get("price", 0))
        qty = int(pos.get("qty", 1))
        summa = price * qty
        grand_total += summa

        # Копируем образец СТРОКИ ДАННЫХ (5 ячеек), а не итога
        new_row = copy.deepcopy(sample_row_xml)
        total_row._element.addprevious(new_row)

        row_obj = _Row(new_row, table)
        cells = row_obj.cells

        _set_cell_text(cells[0], idx, align="center", size=9)
        _set_name_cell(cells[1], pos.get("name", ""), size=9)
        _set_cell_text(cells[2], _format_price(price), align="center", size=9)
        _set_cell_text(cells[3], qty, align="center", size=9)
        _set_cell_text(cells[4], _format_price(summa), align="center", size=9)

    # Обновляем итоговую строку (у неё 2 ячейки: "Итого" gridSpan=4 и сумма)
    total_cells = total_row.cells
    _set_cell_text(total_cells[0], "Итого:", bold=True, size=9)
    # Последняя ячейка — сумма (индекс -1 надёжнее чем 4)
    _set_cell_text(total_cells[-1], _format_price(grand_total), bold=True, align="center", size=9)

    doc.save(output_path)
    logger.info("КП сохранён: %s (итого %.2f)", output_path, grand_total)
    return output_path


if __name__ == "__main__":
    # Тест
    test_positions = [
        {"name": "Распределительное устройство высокого напряжения РУВН (с кабельными перемычками РУ-10кВ - Трансформаторы), в составе ячеек серии КСО-СК-312", "price": 1049911.00, "qty": 1},
        {"name": "Распределительное устройство низкого напряжения РУНН на базе НКУ, с автоматическими выключателями ТЕНГЕН ЭЛЕКТРИК", "price": 2991213.00, "qty": 1},
        {"name": "Трансформатор силовой Сухой ТСЛ 1250 кВА 6/0,4 кВ D/Yн-11", "price": 1776600.00, "qty": 2},
    ]
    generate_kp(
        output_path="/home/claude/test_kp.docx",
        object_name="ООО «Тестовый объект» г. Барнаул",
        equipment_title="Комплектная трансформаторная подстанция 2КТПВ-СК-КК-1250кВА-6/0,4кВ УХЛ1",
        positions=test_positions,
        delivery_address="г. Барнаул, ул. Тестовая, 1",
    )
    print("Test KP generated")